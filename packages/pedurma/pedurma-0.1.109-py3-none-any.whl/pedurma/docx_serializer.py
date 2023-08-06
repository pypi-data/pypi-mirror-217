import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from pypandoc import convert_text


def split_text(content):

    chunks = re.split(r"(\(\d+\) <.+?>)", content)

    return chunks


def create_docx_with_footnotes(text_id, collated_text, path):
    collated_text = collated_text.replace("1-100000", "")
    chunks = split_text(collated_text)
    document = Document()
    p = document.add_paragraph()

    for chunk in chunks:
        if chunk and re.search(r"\(\d+\) <.+?>", chunk):
            chunk = re.sub(r"\(\d+\) ", "", chunk)
            super_text = p.add_run(chunk)
            super_text.font.superscript = True
            super_text.font.name = "Jomolhari"
        else:
            normal_text = p.add_run(chunk)
            normal_text.font.name = "Jomolhari"
    output_path = path / f"{text_id}_format_01.docx"
    document.save(str(output_path))
    return output_path


def get_pages(text):
    result = []
    pg_text = ""
    pages = re.split(r"(\d+-\d+)", text)
    for i, page in enumerate(pages[:-1]):
        if i % 2 == 0:
            pg_text += page
        else:
            pg_text += page
            result.append(pg_text)
            pg_text = ""
    return result


def parse_page(page, note_walker):
    page_ann = re.search(r"(\d+-\d+)", page)[0]
    page = page.replace(page_ann, "")
    page_md = ""
    chunks = split_text(page)
    for chunk in chunks:
        if chunk and re.search(r"\(\d+\) <.+?>", chunk):
            page_md += f"[^{note_walker}]"
            note_walker += 1
        else:
            page_md += chunk
    page_md += f"\n{page_ann}\n"
    return page_md, note_walker


def reformat_namgyal_mon_format(notes):
    reformated_note_text = ""
    for pub, note in notes.items():
        reformated_note_text += f"{note} {pub}"
    full_names = {
        "«སྡེ་»": "སྡེ་དགེ།",
        "«ཅོ་»": "ཅོ་ནེ།",
        "«པེ་»": "པེ་ཅིན།",
        "«སྣར་»": "སྣར་ཐང་།",
    }
    for tib_abv, full_name in full_names.items():
        reformated_note_text = reformated_note_text.replace(tib_abv, f" {full_name} ")
    return reformated_note_text


def reformat_kumarajiva_format(notes):
    reformated_note_text = ""
    for note_walker, (pub, note) in enumerate(notes.items()):
        pub = re.sub("»«", "»,«", pub)
        if len(notes) > 1 and note_walker >= 0 and note_walker < len(notes) - 1:
            reformated_note_text += f"{pub}: {note};"
        else:
            reformated_note_text += f"{pub}: {note}"

    pub_abv = {"«པེ་»": "P", "«སྣར་»": "N", "«ཅོ་»": "C", "«སྡེ་»": "D"}
    for tib_abv, eng_abv in pub_abv.items():
        reformated_note_text = reformated_note_text.replace(tib_abv, f"{eng_abv}")
    return reformated_note_text


def reformat_note_text(note_text, lang="bo"):
    reformated_note_text = ""
    note_parts = re.split("(«.+?»)", note_text)
    notes = {}
    cur_pub = ""
    for note_part in note_parts[1:]:
        if note_part:
            if "«" in note_part:
                cur_pub += note_part
            else:
                notes[cur_pub] = note_part
                cur_pub = ""
    if lang == "bo":
        reformated_note_text = reformat_namgyal_mon_format(notes)
    else:
        reformated_note_text = reformat_kumarajiva_format(notes)
    return reformated_note_text


def reformat_title_note_text(note_text, lang):
    """Reformat the title note text

    Args:
        note_text (str): note text
        lang (str): languange code

    Returns:
        str: reformated title note text
    """
    reformated_note_text = note_text
    if lang == "bo":
        abv_replacement = {
            "«སྡེ་»": "སྡེ་དགེ།",
            "«ཅོ་»": "ཅོ་ནེ།",
            "«པེ་»": "པེ་ཅིན།",
            "«སྣར་»": "སྣར་ཐང་།",
        }
    else:
        abv_replacement = {"«པེ་»": "P", "«སྣར་»": "N", "«ཅོ་»": "C", "«སྡེ་»": "D"}
    for abv, abv_alt in abv_replacement.items():
        reformated_note_text = reformated_note_text.replace(abv, f"{abv_alt}")
    return reformated_note_text


def parse_note(collated_text, lang):
    note_md = "\n"
    notes = re.finditer(r"\((\d+)\) <(.+?)>", collated_text)
    for note_walker, note in enumerate(notes, 1):
        if note_walker == 1:
            note_text = reformat_title_note_text(note.group(2), lang)
        else:
            note_text = reformat_note_text(note.group(2), lang)
        note_md += f"[^{note_walker}]: {note_text}\n"
    return note_md


def creat_docx_footnotes_at_end_of_page(text_id, collated_text, path):
    collated_text_md_nam = ""
    collated_text_md_kuma = ""
    note_walker = 1
    pages = get_pages(collated_text)
    for page in pages:
        page_md, note_walker = parse_page(page, note_walker)
        collated_text_md_kuma += page_md
    collated_text_md_kuma = collated_text_md_kuma.replace("1-100000", "")
    collated_text_md_nam = collated_text_md_kuma
    collated_text_md_kuma += parse_note(collated_text, lang="en")
    collated_text_md_nam += parse_note(collated_text, lang="bo")
    output_path_nam = path / f"{text_id}_format_namgyal.docx"
    output_path_kuma = path / f"{text_id}_format_kumarajiva.docx"
    convert_text(
        collated_text_md_nam, "docx", "markdown", outputfile=str(output_path_nam)
    )
    convert_text(
        collated_text_md_kuma, "docx", "markdown", outputfile=str(output_path_kuma)
    )
    return output_path_kuma


def get_docx_text(text_id, preview_text, output_path=None, type_="with_footnotes"):
    if not output_path:
        (Path.home() / ".collation_docx").mkdir(parents=True, exist_ok=True)
        output_path = Path.home() / ".collation_docx"
    collated_text = ""
    for vol_id, text in preview_text.items():
        collated_text += f"{text}\n\n"
    collated_text = collated_text.replace("\n", "")
    collated_text = re.sub(r"(\d+-\d+)", r"\n\g<1>\n", collated_text)
    if type_ == "with_footnotes":
        docx_path = create_docx_with_footnotes(text_id, collated_text, output_path)
    else:
        docx_path = creat_docx_footnotes_at_end_of_page(
            text_id, collated_text, output_path
        )
    return docx_path
